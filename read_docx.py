from docx import Document
import sys

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(' | '.join(row_text))
    return '\n'.join(full_text)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python read_docx.py <docx_file>")
        sys.exit(1)
    content = read_docx(sys.argv[1])
    print(content)
